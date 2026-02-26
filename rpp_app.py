import streamlit as st
import google.generativeai as genai
from io import BytesIO
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup

# ==============================
# 1. KONFIGURASI HALAMAN & TEMA HITAM
# ==============================
st.set_page_config(page_title="RPP Generator Pro - Andy Kurniawan", page_icon="📄", layout="wide")

st.markdown("""
    <style>
    /* Latar belakang keseluruhan hitam */
    .stApp {
        background-color: #000000;
        color: #ffffff;
    }
    
    /* Header Utama */
    .main-title {
        color: #ffffff;
        text-align: center;
        padding: 30px;
        background: #1e1e1e;
        border-radius: 15px;
        border: 1px solid #3b82f6;
        margin-bottom: 25px;
    }

    /* Form & Input area */
    [data-testid="stForm"] {
        background-color: #111111;
        padding: 30px;
        border-radius: 15px;
        border: 1px solid #444444;
    }

    /* Teks label agar putih */
    label, .stMarkdown p, .stSelectbox label {
        color: #ffffff !important;
        font-weight: bold;
    }

    /* Kartu per pertemuan */
    .meeting-card {
        background-color: #1e3a8a;
        padding: 10px 15px;
        border-radius: 8px;
        margin-top: 20px;
        color: white;
        font-weight: bold;
    }

    /* Input text box & Select box */
    input, textarea, [data-baseweb="select"] {
        background-color: #222222 !important;
        color: white !important;
    }
    
    /* Hasil RPP Preview Box (Kertas Putih) */
    .preview-box {
        background-color: #ffffff;
        color: #000000;
        padding: 40px;
        border-radius: 10px;
        margin-top: 20px;
        border: 1px solid #ccc;
    }
    </style>
    """, unsafe_allow_html=True)

# ==============================
# 2. SECRETS & API
# ==============================
GEMINI_API_KEY = st.secrets.get("GEMINI_API_KEY", "")
APP_PASSWORD = st.secrets.get("APP_PASSWORD", "12345")

# ==============================
# 3. LOGIN SYSTEM
# ==============================
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.markdown("<div class='main-title'><h1>🔐 Akses Guru RPP Pro</h1><p>Gunakan password untuk masuk</p></div>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1,2,1])
    with c2:
        pwd = st.text_input("Password", type="password")
        if st.button("Masuk"):
            if pwd == APP_PASSWORD:
                st.session_state.authenticated = True
                st.rerun()
            else:
                st.error("❌ Password salah")
    st.stop()

# ==============================
# 4. FUNGSI EKSPOR WORD
# ==============================
def create_word(html_content):
    soup = BeautifulSoup(html_content, "html.parser")
    doc = Document()
    
    # Set Font Default
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for el in soup.find_all(["h1", "h2", "h3", "p", "table"]):
        if el.name in ["h1", "h2", "h3"]:
            doc.add_heading(el.get_text(), level=2)
        elif el.name == "p":
            doc.add_paragraph(el.get_text())
        elif el.name == "table":
            rows = el.find_all("tr")
            if not rows: continue
            
            # Hitung kolom maksimal
            max_cols = 0
            for r in rows:
                max_cols = max(max_cols, len(r.find_all(["td", "th"])))
            
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = "Table Grid"
            
            for i, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for j, cell in enumerate(cells):
                    if j < max_cols:
                        table.cell(i, j).text = cell.get_text(strip=True)
                        
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==============================
# 5. FUNGSI GENERATE RPP (STABIL)
# ==============================
def generate_rpp(data):
    if not GEMINI_API_KEY:
        return "<p style='color:red;'>Error: API KEY tidak ditemukan di Secrets!</p>"
    
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        # Menggunakan nama model yang paling standar
        model = genai.GenerativeModel(model_name='gemini-1.5-flash')
        
        pertemuan_str = "\n".join([f"P{i+1}: Model {p['model']}, Waktu {p['waktu']}, Tanggal {p['tanggal']}" for i, p in enumerate(data['pertemuan'])])
        
        prompt = f"""
        Buatlah RPP Kurikulum Merdeka Lengkap dalam format HTML (gunakan tag table, tr, td, th dengan border="1").
        IDENTITAS:
        Sekolah: {data['sekolah']}, Mapel: {data['mapel']}, Fase/Kelas: {data['fase']}.
        MATERI UTAMA: {data['materi']}
        TUJUAN PEMBELAJARAN: {data['tujuan']}
        
        RINCIAN TIAP PERTEMUAN:
        {pertemuan_str}
        
        SYARAT KONTEN:
        1. Tiap pertemuan wajib ada tabel kegiatan: LANGKAH | DESKRIPSI KEGIATAN | WAKTU.
        2. Gunakan sintaks yang sesuai dengan model pembelajaran yang dipilih.
        3. Tambahkan bagian Asesmen (Diagnostik, Formatif, Sumatif) dan Rubrik Nilai.
        4. Berikan penutup dan tanda tangan: 
           Mengetahui, Kepala Sekolah (AHMAD JUNAIDI, S.Pd) dan Guru Mata Pelajaran (ANDY KURNIAWAN, S.Pd.SD).
        """
        
        response = model.generate_content(prompt)
        # Bersihkan tag markdown jika ada
        return response.text.replace("```html", "").replace("```", "").strip()
    except Exception as e:
        return f"<p style='color:red;'>Terjadi kesalahan API Google: {str(e)}</p>"

# ==============================
# 6. DAFTAR MODEL PEMBELAJARAN
# ==============================
MODELS_LIST = [
    "Problem Based Learning (PBL)", "Project Based Learning (PjBL)", 
    "Discovery Learning", "Inquiry Learning", "Contextual Teaching (CTL)",
    "Cooperative Learning (Jigsaw)", "Cooperative Learning (STAD)",
    "Demonstrasi & Eksperimen", "Direct Instruction (Ceramah)", 
    "Role Playing (Bermain Peran)", "Numbered Head Together (NHT)", 
    "Mind Mapping", "Talking Stick", "Game Based Learning (PJOK/Umum)", 
    "Inside Outside Circle", "Example Non-Example", "Think Pair Share (TPS)",
    "Problem Solving (PAI/Agama)", "Kisah Keteladanan (Agama)", "Praktik Lapangan (Olahraga)"
]

# ==============================
# 7. UI UTAMA
# ==============================
st.markdown("<div class='main-title'><h1>📄 RPP GENERATOR PRO</h1><p>Andalan Guru Indonesia - By Andy Kurniawan, S.Pd.SD</p></div>", unsafe_allow_html=True)

with st.form("rpp_form"):
    st.subheader("📋 Identitas RPP")
    c1, c2 = st.columns(2)
    with c1:
        sekolah = st.text_input("Nama Sekolah", "SDN ...")
        mapel = st.text_input("Mata Pelajaran", "PJOK / PAI / Guru Kelas")
    with c2:
        fase = st.text_input("Fase / Kelas / Semester", "B / IV / 1")
        jml = st.number_input("Jumlah Pertemuan (Maks 15)", 1, 15, 1)
    
    st.divider()
    
    st.subheader("📅 Detail Tiap Pertemuan")
    pertemuan_data = []
    for i in range(int(jml)):
        st.markdown(f"<div class='meeting-card'>Pertemuan Ke-{i+1}</div>", unsafe_allow_html=True)
        col1, col2, col3 = st.columns([2, 1, 1])
        with col1:
            m = st.selectbox(f"Model P{i+1}", MODELS_LIST, key=f"m{i}")
        with col2:
            w = st.text_input(f"Waktu P{i+1}", "2 x 35 Menit", key=f"w{i}")
        with col3:
            t = st.date_input(f"Tanggal P{i+1}", key=f"t{i}")
        pertemuan_data.append({"model": m, "waktu": w, "tanggal": t.strftime("%d/%m/%Y")})
    
    st.divider()
    
    st.subheader("📝 Konten Pembelajaran")
    materi = st.text_area("Ringkasan Materi Utama", height=100, placeholder="Contoh: Variasi gerak dasar lokomotor...")
    tujuan = st.text_area("Tujuan Pembelajaran (TP)", height=100, placeholder="Contoh: Peserta didik mampu mempraktikkan...")
    
    submit = st.form_submit_button("🚀 GENERATE RPP SEKARANG")

# ==============================
# 8. OUTPUT & EKSPOR
# ==============================
if submit:
    if not materi or not tujuan:
        st.error("⚠️ Materi dan Tujuan Pembelajaran tidak boleh kosong!")
    else:
        with st.spinner("🧠 AI sedang menyusun RPP lengkap... Mohon tunggu."):
            data_input = {
                "sekolah": sekolah, "mapel": mapel, "fase": fase, 
                "materi": materi, "tujuan": tujuan, "pertemuan": pertemuan_data
            }
            hasil_html = generate_rpp(data_input)
            
            if "Terjadi kesalahan" in hasil_html:
                st.error(hasil_html)
            else:
                st.success("✅ RPP Selesai Dibuat!")
                
                # Tampilkan Preview
                st.markdown("<div class='preview-box'>", unsafe_allow_html=True)
                st.html(hasil_html)
                st.markdown("</div>", unsafe_allow_html=True)
                
                # Tombol Download Word
                file_docx = create_word(hasil_html)
                st.download_button(
                    label="📥 Download File Word (.docx)",
                    data=file_docx,
                    file_name=f"RPP_{mapel}_{fase}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

st.markdown("<br><p style='text-align: center; color: #555;'>© 2026 RPP Generator Pro | Andy Kurniawan, S.Pd.SD</p>", unsafe_allow_html=True)
